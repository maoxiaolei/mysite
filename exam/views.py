#_*_coding:utf-8_*_

from django.shortcuts import render
from django.http import HttpResponse,HttpResponseRedirect
from django.views import generic
from .models import Question
from django.urls import reverse
from docx import Document
from django.utils import timezone
import time
import os
from django.contrib.auth.decorators import login_required


class IndexView(generic.ListView):
    template_name = 'exam/index.html'
    context_object_name = 'question_list'
    def get_queryset(self):
        return Question.objects.all()

class DetailView(generic.DetailView):
    model = Question
    template_name = 'exam/detail.html'


@login_required()
def generate(request):
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph(u'加点中文，可以吗')
    for question in Question.objects.all():
        document.add_paragraph(question.question_text)
        for choice in question.choice_set.all():
            if choice <> '':
                document.add_paragraph(choice.choice_text)
        for content in question.content_set.all():
            if content <> '':
                document.add_paragraph(content.content_text)
        document.add_paragraph(' ')

    document.add_page_break()

    document.save('exam/static/demo'+time.strftime('%Y-%m-%d %H%M%S',time.localtime(time.time()))+'.docx')

    rootdir = 'exam/static/'
    filenames = os.walk(rootdir)
    return render(request, 'exam/generate.html', {
        'filenames': filenames,
    })

